#!/usr/bin/env python3
"""
eecc_server.py — Servidor para generación de EECC
n8n llama a POST /generar con URLs de archivos y parámetros del cliente.
"""
import os, re, io, subprocess, tempfile, shutil, zipfile, urllib.request
from pathlib import Path
from datetime import date, datetime
from fastapi import FastAPI, Form, HTTPException
from fastapi.responses import FileResponse
import uvicorn

app = FastAPI(title="EECC Generator")

GEN_SCRIPT       = Path(__file__).parent / "gen_eecc_v4.py"
INFORME_TEMPLATE = Path(__file__).parent / "informe_template.docx"

MONTHS_ES = {
    1:'enero', 2:'febrero', 3:'marzo', 4:'abril', 5:'mayo', 6:'junio',
    7:'julio', 8:'agosto', 9:'septiembre', 10:'octubre', 11:'noviembre', 12:'diciembre'
}


@app.get("/health")
def health():
    import shutil, subprocess as sp
    lo = shutil.which("libreoffice") or shutil.which("soffice") or "NOT FOUND"
    try:
        ver = sp.run([lo, "--version"], capture_output=True, text=True, timeout=10).stdout.strip()
    except Exception as e:
        ver = str(e)
    return {
        "status": "ok",
        "script": str(GEN_SCRIPT),
        "script_exists": GEN_SCRIPT.exists(),
        "template_exists": INFORME_TEMPLATE.exists(),
        "libreoffice_path": lo,
        "libreoffice_version": ver,
    }


@app.post("/generar")
async def generar(
    ss_url:        str   = Form(...),
    eecc_url:      str   = Form(default=""),
    empresa:       str   = Form(...),
    cuit:          str   = Form(...),
    domicilio:     str   = Form(default=""),
    matricula_igj: str   = Form(default=""),
    nro_ejercicio: int   = Form(default=1),
    fecha_cierre:  str   = Form(...),   # YYYY-MM-DD
    cof:           float = Form(...),
    cap_nominal:   float = Form(...),
    sipa_monto:    str   = Form(default=""),
):
    tmp = tempfile.mkdtemp(prefix="eecc_")
    try:
        ss_act_path = os.path.join(tmp, "ss_actual.xlsx")
        out_path    = os.path.join(tmp, "output.xlsx")

        urllib.request.urlretrieve(ss_url, ss_act_path)

        cmd = [
            "python3", str(GEN_SCRIPT),
            "--empresa",       empresa,
            "--cuit",          cuit,
            "--nro-ejercicio", str(nro_ejercicio),
            "--fecha-cierre",  fecha_cierre.strip(),
            "--cof",           str(cof),
            "--cap-nominal",   str(cap_nominal),
            "--ss-actual",     ss_act_path,
            "--output",        out_path,
        ]

        if eecc_url and eecc_url.strip():
            prev_path = os.path.join(tmp, "eecc_anterior.pdf")
            urllib.request.urlretrieve(eecc_url.strip(), prev_path)
            if os.path.getsize(prev_path) > 100:
                cmd += ["--eecc-anterior", prev_path]

        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if result.returncode != 0:
            raise HTTPException(status_code=500,
                detail=f"Error en gen_eecc: {result.stderr}\n{result.stdout}")

        if not os.path.exists(out_path):
            raise HTTPException(status_code=500, detail="El script no generó el archivo")

        empresa_slug = empresa.replace(" ", "_").replace(".", "")[:30]
        cuit_slug    = cuit.replace("-", "").replace(" ", "")
        year         = fecha_cierre.strip()[:4]
        xlsx_name    = f"EECC_{cuit_slug}_{year}.xlsx"
        pdf_name     = f"EECC_{cuit_slug}_{year}.pdf"
        excel_pdf    = os.path.join(tmp, "excel.pdf")
        informe_pdf  = os.path.join(tmp, "informe.pdf")
        merged_pdf   = os.path.join(tmp, pdf_name)

        # 1. Excel → PDF
        cc_key = os.environ.get("CLOUDCONVERT_API_KEY", "")
        lo = _find_libreoffice()
        if cc_key:
            try:
                _cloudconvert_pdf(cc_key, out_path, excel_pdf)
            except Exception as cc_err:
                print(f"[CloudConvert] falló ({cc_err}), usando fallback")
                if lo:
                    _libreoffice_convert(lo, out_path, tmp, excel_pdf)
                else:
                    _xlsx_to_pdf(out_path, excel_pdf)
        elif lo:
            _libreoffice_convert(lo, out_path, tmp, excel_pdf)
        else:
            _xlsx_to_pdf(out_path, excel_pdf)

        # 2. Informe Word → rellenar → PDF
        if INFORME_TEMPLATE.exists():
            informe_filled = os.path.join(tmp, "informe_filled.docx")
            _fill_informe(str(INFORME_TEMPLATE), informe_filled,
                          empresa, cuit, domicilio, matricula_igj,
                          fecha_cierre.strip(), sipa_monto)
            if lo:
                _libreoffice_convert(lo, informe_filled, tmp, informe_pdf)
            else:
                _docx_to_pdf(informe_filled, informe_pdf)

        # 2b. Notas contables → docx → PDF
        notas_pdf = os.path.join(tmp, "notas.pdf")
        notes_json_path = os.path.join(tmp, "notes_data.json")
        if os.path.exists(notes_json_path):
            try:
                import json as _json
                with open(notes_json_path, "r", encoding="utf-8") as _f:
                    notes = _json.load(_f)
                notas_docx = os.path.join(tmp, "notas.docx")
                _generate_notas_docx(notes, notas_docx, fecha_cierre.strip())
                if lo:
                    _libreoffice_convert(lo, notas_docx, tmp, notas_pdf)
                else:
                    _docx_to_pdf(notas_docx, notas_pdf)
            except Exception as notas_err:
                print(f"[NOTAS] Error generando notas.docx: {notas_err}")
                notas_pdf = None
        else:
            notas_pdf = None

        # 3. Mergear PDFs
        pdfs_to_merge = [excel_pdf]
        if notas_pdf and os.path.exists(notas_pdf):
            pdfs_to_merge.append(notas_pdf)
        if os.path.exists(informe_pdf):
            pdfs_to_merge.append(informe_pdf)
        _merge_pdfs(pdfs_to_merge, merged_pdf)

        # 4. ZIP con xlsx + pdf
        zip_path = os.path.join(tmp, "eecc.zip")
        with zipfile.ZipFile(zip_path, "w") as zf:
            zf.write(out_path, xlsx_name)
            if os.path.exists(merged_pdf):
                zf.write(merged_pdf, pdf_name)

        final_zip = Path(tempfile.gettempdir()) / f"eecc_{os.path.basename(tmp)}.zip"
        shutil.copy(zip_path, final_zip)

        return FileResponse(
            path=str(final_zip),
            media_type="application/zip",
            filename="eecc.zip",
            background=_cleanup(tmp, final_zip),
        )

    except HTTPException:
        shutil.rmtree(tmp, ignore_errors=True)
        raise
    except Exception as e:
        shutil.rmtree(tmp, ignore_errors=True)
        raise HTTPException(status_code=500, detail=str(e))


def _cloudconvert_pdf(api_key: str, input_path: str, output_path: str):
    """Convierte a PDF via CloudConvert API (usa LibreOffice internamente)."""
    import cloudconvert
    cloudconvert.configure(api_key=api_key, sandbox=False)

    job = cloudconvert.Job.create(payload={
        "tasks": {
            "upload":  {"operation": "import/upload"},
            "convert": {"operation": "convert", "input": "upload",
                        "output_format": "pdf", "engine": "libreoffice"},
            "export":  {"operation": "export/url", "input": "convert"}
        }
    })

    upload_task = next(t for t in job["tasks"] if t["name"] == "upload")
    cloudconvert.Task.upload(file_name=input_path, task=upload_task)

    job = cloudconvert.Job.wait(id=job["id"])
    export_task = next(t for t in job["tasks"] if t["name"] == "export")
    url = export_task["result"]["files"][0]["url"]
    urllib.request.urlretrieve(url, output_path)


def _find_libreoffice():
    import shutil
    return shutil.which("libreoffice") or shutil.which("soffice")


def _libreoffice_convert(lo_bin: str, input_path: str, out_dir: str, desired_path: str):
    """Convierte un archivo a PDF con LibreOffice headless."""
    env = os.environ.copy()
    env["HOME"] = out_dir  # evita conflictos de perfil de usuario
    subprocess.run(
        [lo_bin, "--headless", "--norestore", "--convert-to", "pdf",
         "--outdir", out_dir, input_path],
        capture_output=True, timeout=120, env=env
    )
    base = os.path.splitext(os.path.basename(input_path))[0]
    generated = os.path.join(out_dir, base + ".pdf")
    if os.path.exists(generated) and generated != desired_path:
        os.rename(generated, desired_path)


def _xlsx_to_pdf(xlsx_path: str, pdf_path: str):
    """Convierte Excel a PDF: una página por solapa preservando estilos de xlsx2html."""
    from xlsx2html import xlsx2html as x2h
    from weasyprint import HTML
    from openpyxl import load_workbook
    from openpyxl.utils import get_column_letter
    from pypdf import PdfWriter, PdfReader
    import tempfile

    LANDSCAPE_SHEETS = {'EEPN', 'Anexo I', 'Anexo III'}
    MULTIPAGE_OK     = {'Notas'}   # solapas que pueden tener varias páginas
    MARGIN_CM        = 0.5         # margen en cm para hojas financieras

    def _sheet_scale(ws, landscape):
        """Calcula el factor de zoom para que la hoja entre en 1 página A4."""
        page_w_mm = 297 if landscape else 210
        page_h_mm = 210 if landscape else 297
        usable_w_px = (page_w_mm - 2 * MARGIN_CM * 10) / 25.4 * 96
        usable_h_px = (page_h_mm - 2 * MARGIN_CM * 10) / 25.4 * 96

        n_cols = ws.max_column or 1
        n_rows = ws.max_row or 1

        # Ancho total: xlsx2html usa ~7.2px por unidad Excel; +padding
        col_w = sum(
            (ws.column_dimensions.get(get_column_letter(c)) or
             type('_', (), {'width': 8})()).width
            for c in range(1, n_cols + 1)
        )
        # Alto total: 1 punto Excel ≈ 1.333px
        row_h = sum(
            (ws.row_dimensions.get(r) or
             type('_', (), {'height': 15})()).height
            for r in range(1, n_rows + 1)
        )

        content_w_px = col_w  * 7.2 + 30   # +30 para bordes/scroll
        content_h_px = row_h * 1.333 + 30

        scale_w = usable_w_px / max(content_w_px, 1)
        scale_h = usable_h_px / max(content_h_px, 1)
        return min(1.0, scale_w, scale_h)

    wb = load_workbook(xlsx_path)
    sheet_pdfs = []

    for sheet_name in wb.sheetnames:
        buf = io.StringIO()
        try:
            x2h(xlsx_path, buf, sheet=sheet_name)
            full_html = buf.getvalue()

            landscape = sheet_name in LANDSCAPE_SHEETS
            multipage = sheet_name in MULTIPAGE_OK
            pw = '297mm' if landscape else '210mm'
            ph = '210mm' if landscape else '297mm'

            ws_sheet = wb[sheet_name]

            if multipage:
                inject = (
                    f'<style>'
                    f'@page {{ size: {pw} {ph}; margin: 0.7cm; }}'
                    f'body {{ font-size: 7pt !important; }}'
                    f'table {{ width: 100% !important; }}'
                    f'td, th {{ white-space: normal !important; word-break: break-word !important; }}'
                    f'</style>'
                )
            else:
                scale = _sheet_scale(ws_sheet, landscape)
                print(f"[XLSX2PDF] {sheet_name}: scale={scale:.3f} ({'landscape' if landscape else 'portrait'})")
                inject = (
                    f'<style>'
                    f'@page {{ size: {pw} {ph}; margin: {MARGIN_CM}cm; }}'
                    # zoom escala tanto ancho como alto colapsando el espacio en layout
                    f'html {{ zoom: {scale:.4f}; }}'
                    f'table {{ border-collapse: collapse !important; }}'
                    f'td, th {{ overflow: hidden !important; white-space: nowrap !important; '
                    f'padding: 0 2px !important; line-height: 1.3 !important; }}'
                    f'</style>'
                )

            if '</head>' in full_html:
                full_html = full_html.replace('</head>', inject + '</head>')
            else:
                full_html = inject + full_html

            tmp = tempfile.mktemp(suffix='.pdf')
            HTML(string=full_html).write_pdf(tmp)
            sheet_pdfs.append(tmp)
        except Exception as e:
            print(f"[XLSX2PDF] {sheet_name}: ERROR {e}")

    writer = PdfWriter()
    for p in sheet_pdfs:
        if os.path.exists(p):
            for page in PdfReader(p).pages:
                writer.add_page(page)
            os.unlink(p)

    with open(pdf_path, 'wb') as f:
        writer.write(f)


def _docx_to_pdf(docx_path: str, pdf_path: str):
    """Convierte DOCX a PDF preservando alineación e imágenes."""
    import base64
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from weasyprint import HTML

    doc = Document(docx_path)
    NS_DRAW = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
    NS_A    = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
    NS_R    = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

    def _para_html(para):
        align_map = {
            WD_ALIGN_PARAGRAPH.RIGHT:   'right',
            WD_ALIGN_PARAGRAPH.CENTER:  'center',
            WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
            WD_ALIGN_PARAGRAPH.LEFT:    'left',
        }
        align = align_map.get(para.alignment, 'justify')
        parts = []
        for run in para.runs:
            inline = run._element.find(f'.//{NS_DRAW}inline')
            if inline is not None:
                blip = run._element.find(f'.//{NS_A}blip')
                if blip is not None:
                    rId = blip.get(f'{{{NS_R}}}embed')
                    if rId:
                        img_part = doc.part.related_parts[rId]
                        b64 = base64.b64encode(img_part.blob).decode()
                        mime = img_part.content_type
                        parts.append(f'<img src="data:{mime};base64,{b64}" style="width:7.1cm;" />')
            else:
                text = run.text.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
                if not text:
                    continue
                s = ''
                if run.bold:    s += 'font-weight:bold;'
                if run.italic:  s += 'font-style:italic;'
                if run.underline: s += 'text-decoration:underline;'
                parts.append(f'<span style="{s}">{text}</span>' if s else text)
        content = ''.join(parts) if parts else '&nbsp;'
        return f'<p style="text-align:{align};margin:0.25em 0">{content}</p>'

    html_parts = [_para_html(p) for p in doc.paragraphs]

    css = '''
        @page { size: A4; margin: 2.5cm 2cm; }
        body { font-family: Arial, sans-serif; font-size: 11pt; line-height: 1.5; }
    '''
    full_html = f'<html><head><style>{css}</style></head><body>{"".join(html_parts)}</body></html>'
    HTML(string=full_html).write_pdf(pdf_path)


def _fill_informe(template_path: str, out_path: str,
                  empresa: str, cuit: str, domicilio: str, matricula_igj: str,
                  fecha_cierre: str, sipa_monto: str):
    """Rellena el template del Informe de Auditoría con los datos del cliente."""
    from docx import Document

    fecha = datetime.strptime(fecha_cierre, "%Y-%m-%d")
    fecha_larga  = f"{fecha.day} de {MONTHS_ES[fecha.month]} de {fecha.year}"
    mes_anio     = f"{MONTHS_ES[fecha.month]} de {fecha.year}"
    today        = date.today()
    fecha_inf    = f"{today.day} de {MONTHS_ES[today.month]} de {today.year}"
    sipa_fmt     = sipa_monto.strip() if sipa_monto.strip() else "[COMPLETAR SIPA]"

    replacements = {
        "{{EMPRESA}}.": f"{empresa}.",
        "{{EMPRESA}}":  empresa,
        "{{CUIT}}":     cuit,
        "{{DOMICILIO}}":          domicilio or "[DOMICILIO]",
        "{{MATRICULA_IGJ}}":      matricula_igj or "[MATRÍCULA]",
        "{{FECHA_CIERRE_LARGA}}": fecha_larga,
        "{{MES_ANIO_CIERRE}}":    mes_anio,
        "{{SIPA_MONTO}}":         sipa_fmt,
        "{{FECHA_INFORME}}":      fecha_inf,
    }

    doc = Document(template_path)
    for para in doc.paragraphs:
        _replace_para(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_para(para, replacements)
    doc.save(out_path)


def _replace_para(para, replacements):
    for key, val in replacements.items():
        if key in para.text:
            for run in para.runs:
                if key in run.text:
                    run.text = run.text.replace(key, val)


def _generate_notas_docx(notes: dict, out_path: str, fecha_cierre: str):
    """Genera el documento Word de Notas a los Estados Contables."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── Márgenes de página ──────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Helpers ─────────────────────────────────────────────────────────────
    def _add_para(text, bold=False, italic=False, size=10,
                  align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
        p = doc.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after  = Pt(space_after)
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        return p

    def _fmt(v):
        """Formatea un número como $ 1,234,567.89"""
        if v is None:
            return "$ -"
        try:
            return f"$ {float(v):>15,.2f}"
        except (TypeError, ValueError):
            return str(v)

    def _set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _add_table(label, rows_data):
        """
        rows_data: list of (cuenta, val25, val24)
        Última fila = TOTAL (bold).
        """
        _add_para(label, bold=True, size=10, space_before=8, space_after=2)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"

        # Anchos de columna
        tbl.columns[0].width = Cm(10)
        tbl.columns[1].width = Cm(3.5)
        tbl.columns[2].width = Cm(3.5)

        # Header
        hdr = tbl.rows[0].cells
        hdr[0].text = "Cuenta"
        hdr[1].text = f"Actual\n{notes['ej25']}"
        hdr[2].text = f"Anterior (reexp.)\n{notes['ej24']}"
        for cell in hdr:
            _set_cell_bg(cell, "EEEEEE")
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        # Filas de datos
        for i, (cuenta, v25, v24) in enumerate(rows_data):
            is_total = (i == len(rows_data) - 1)
            row = tbl.add_row().cells
            row[0].text = cuenta
            row[1].text = _fmt(v25)
            row[2].text = _fmt(v24)
            for j, cell in enumerate(row):
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in para.runs:
                        run.bold = is_total
                        run.font.size = Pt(9)

        doc.add_paragraph()  # espacio después de tabla

    # ── Valores de notas ─────────────────────────────────────────────────────
    n = notes
    ej25      = n["ej25"]
    ej24      = n["ej24"]
    ej25_year = ej25.split("/")[-1] if "/" in ej25 else ej25[:4]

    # ── HEADER ───────────────────────────────────────────────────────────────
    _add_para(n["empresa"], bold=True, size=14,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _add_para("NOTAS A LOS ESTADOS CONTABLES", bold=True, size=12,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _add_para(f"Por el ejercicio cerrado el {ej25}", italic=True, size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    # Línea separadora
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_after = Pt(6)
    p_run = p_sep.add_run("─" * 80)
    p_run.font.size = Pt(8)

    # ── SECCIÓN 1: Texto normativo ────────────────────────────────────────────
    sec1_blocks = [
        ("1 - Normas Contables aplicadas:", True),
        ("A continuación se detallan las normas contables más relevantes, utilizadas por la Sociedad para la preparación de los presentes estados contables.", False),
        ("1.1 Modelo de presentación de los Estados Contables", True),
        (f"Los presentes estados contables han sido preparados en moneda homogénea (pesos de diciembre de {ej25_year}), reconociendo en forma integral los efectos de la inflación de conformidad con lo establecido en la Resolución Tecnica (RT) N° 6, en virtud de haberse determinado la existencia de un contexto de alta inflación que vuelve necesaria la reexpresion de los estados contables.", False),
        ("1.2 Consideración de los efectos del cambio en el poder adquisitivo de la moneda:", True),
        ("Desde la entrada en vigencia de la RT N° 39 (aprobada por el Consejo Profesional de Ciencias Económicas de la Ciudad Autónoma de Buenos Aires (CPCECABA) mediante Resolución de Consejo Directivo N° 20/2014), que modifico las normas sobre la unidad de medida de la RT N° 17, la necesidad de reexpresar los estados contables para reflejar los cambios en el poder adquisitivo de la moneda viene indicada por la existencia o no de un contexto de inflación tal que lleve a calificar la economía de altamente inflacionaria. A los fines de identificar la existencia de un entorno económico inflacionario, la interpretación N° 8 (aprobada por el CPCECABA mediante Resolución del Consejo Directivo N° 115/2014) brinda una pauta cuantitativa que es condición necesaria para proceder a reexpresar las cifras de los estados contables, dicha pauta consiste en que la tasa acumulada de inflación en tres años, considerando el Indice de Precios Internos al por Mayor (IPIM) elaborado por el Instituto de Estadística y Censos (INDEC), alcance o sobrepace el 100 % entre otros factores.", False),
        ("Durante el primer semestre de 2018, diversos factores macroeconómicos produjeron una aceleración significativa de la inflación, resultando en índices que excedieron el 100 % acumulado en tres años, y en proyecciones de inflación que confirmaron dicha tendencia. Como consecuencia de ello, la Junta de Gobierno de la Federación Argentina de Consejos Profesionales de Ciencias Económicas (FACPCE) emitió la Resolución N° 539/2018 (aprobada por el CPCECABA mediante Resolución de Consejo Directivo N° 107/2018), indicando que se encontraba configurado el contexto de alta inflación y que los estados contables correspondientes a períodos anuales o intermedios cerrados a partir del 1 de julio de 2018 deberán ser ajustados para reflejar los cambios en el poder adquisitivo de la moneda.", False),
        ("La aplicación del proceso de reexpresion establecido en la RT N° 6 permite el reconocimiento de las ganancias y perdidas derivadas del mantenimiento de activos y pasivos expuestos a los cambios del poder adquisitivo de la moneda del estado de resultados.", False),
        (f"El estado contable correspondiente al ejercicio cerrado el {ej25}, se encuentra ajustado por inflación.", False),
        ("1.3 Criterios de Valuación", True),
        ("1.3.1 Los activos y pasivos en moneda nacional están valuados a su valor nominal.", True),
        ("1.3.2 Impuesto a las Ganancias", True),
        ("Las normas contables profesionales vigentes requieren la contabilización del impuesto a las ganancias por el método del impuesto diferido. Este criterio implica el reconocimiento de partidas de activos y de pasivos por impuesto diferido, en los casos que se produzcan diferencias temporarias entre la medición contable y la medición fiscal de los activos y de los pasivos, o cuando existan quebrantos impositivos utilizables para compensar ganancias imponibles de ejercicios futuros.", False),
        ("La Sociedad no presenta diferencias temporarias ni quebrantos impositivos utilizables para compensar ganancias imponibles de ejercicios futuros, por lo tanto, determino el cargo por impuesto a las ganancias mediante la aplicación de la tasa de dicho impuesto sobre el resultado impositivo, el cual coincide con el resultado contable por no haber diferencias temporarias en las valuaciones contables e impositivas de los activos y pasivos.", False),
        ("1.3.3 RECPAM: En el estado de resultado en moneda constante, se exponen en forma conjunta bajo la denominación \"RECPAM\" incluyendo resultados por exposición al cambio en el poder adquisitivo de la moneda los siguientes conceptos: Resultados por tenencia, resultados financieros y resultados por exposición al cambio del poder adquisitivo de la moneda.", True),
        ("2 - Composición de los principales rubros:", True),
    ]

    for text, is_bold in sec1_blocks:
        _add_para(text, bold=is_bold, size=10,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3)

    # ── SECCIÓN 2: Tablas numéricas ───────────────────────────────────────────
    def rx(v):
        return round(float(v) * float(n["cof"]), 2)

    _add_table("2.1 Caja y Bancos", [
        ("Caja",             n["caja25_caja"],  rx(800000)),
        ("Banco Santander",  n["caja25_banco"], rx(36865.50)),
        ("TOTAL",            n["caja25"],       n["caja24"]),
    ])

    _add_table("2.2 Créditos por Ventas en Moneda", [
        ("Deudores por Ventas", n["cv25"], n["cv24"]),
        ("TOTAL",               n["cv25"], n["cv24"]),
    ])

    _add_table("2.3 Otros Créditos en Moneda", [
        ("IVA Saldo Técnico",              n["oc25_iva"], rx(323758.96)),
        ("Saldo Libre Disponibilidad",     n["oc25_sld"], rx(655788.38)),
        ("Impuesto Débitos y Créditos",    n["oc25_dbc"], rx(32934.04)),
        ("Retención Ganancias Sufrida",    n["oc25_ret"], rx(738537.29)),
        ("TOTAL",                          n["oc25"],     n["oc24"]),
    ])

    _add_table("2.4 Bienes de Cambio", [
        ("Bienes de Cambio", n["bc25"], n["bc24"]),
        ("TOTAL",            n["bc25"], n["bc24"]),
    ])

    _add_table("2.5.1 Deudas Comerciales en Moneda", [
        ("Proveedores", n["dc25"], n["dc24"]),
        ("TOTAL",       n["dc25"], n["dc24"]),
    ])

    _add_table("2.5.2 Cargas Fiscales en Moneda", [
        ("IIBB BSAS a pagar", n["df25_bsas"], rx(58067)),
        ("IIBB CABA a pagar", n["df25_caba"], 0),
        ("TOTAL",             n["df25"],      n["df24"]),
    ])

    _add_table("2.5.3 Remuneraciones en Moneda", [
        ("Cargas Sociales a pagar", n["rem25"], n["rem24"]),
        ("TOTAL",                   n["rem25"], n["rem24"]),
    ])

    _add_table("2.5.4 Deudas Sociales en Moneda", [
        ("Cuenta particular socios", n["ds25"], n["ds24"]),
        ("TOTAL",                    n["ds25"], n["ds24"]),
    ])

    _add_table("2.6 RECPAM", [
        ("RECPAM operativo",                    n["recpam25"],     n["recpam24"]),
        ("Ajuste reexpresión apertura (RT6)",   n["recpam_rx_aper"], 0),
        ("TOTAL RECPAM",                        n["recpam25_adj"], n["recpam24"]),
    ])

    # ── SECCIÓN 3: Patrimonio Neto Negativo (condicional) ────────────────────
    if float(n["pn25"]) < 0:
        _add_para("3 - Patrimonio Neto Negativo", bold=True, size=10,
                  space_before=8, space_after=3)
        _add_para(
            f"Al cierre del ejercicio {ej25}, la Sociedad presenta un Patrimonio Neto "
            f"negativo de {_fmt(n['pn25'])}. Esta situación obedece a los resultados "
            f"acumulados del ejercicio y ejercicios anteriores. Los socios se encuentran "
            f"al tanto de esta situación y han comprometido su apoyo financiero para "
            f"normalizar la situación patrimonial de la Sociedad.",
            size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3
        )

    # ── FOOTER en cada página ─────────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = "Las notas forman parte integrante de los estados contables."
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_para.runs:
            run.font.size = Pt(8)
            run.italic = True

    doc.save(out_path)
    print(f"[NOTAS] notas.docx guardado: {out_path}")


def _merge_pdfs(pdf_paths: list, output_path: str):
    """Mergea una lista de PDFs en uno solo."""
    try:
        from pypdf import PdfWriter, PdfReader
        writer = PdfWriter()
        for path in pdf_paths:
            if os.path.exists(path):
                reader = PdfReader(path)
                for page in reader.pages:
                    writer.add_page(page)
        with open(output_path, "wb") as f:
            writer.write(f)
    except Exception as e:
        print(f"[MERGE] Error: {e}")


class _cleanup:
    def __init__(self, tmp_dir, final_zip):
        self._tmp = tmp_dir
        self._zip = final_zip
    def __call__(self, *_):
        shutil.rmtree(self._tmp, ignore_errors=True)
        try: os.unlink(self._zip)
        except: pass


if __name__ == "__main__":
    print("Iniciando servidor EECC en http://localhost:8000")
    uvicorn.run(app, host="0.0.0.0", port=8000)
